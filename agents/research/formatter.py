"""Word document formatter for the weekly research brief.

Matches the Plumtree style: purple headers, clean sections, professional layout.
"""

import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Plumtree brand purple
PLUMTREE_PURPLE = RGBColor(0x6B, 0x2D, 0x8B)
PLUMTREE_DARK = RGBColor(0x33, 0x33, 0x33)
HEADER_BG_PURPLE = RGBColor(0x4A, 0x15, 0x6B)

CATEGORIES = [
    "Teams and Teaming",
    "Organization Systems",
    "Culture",
    "Leadership",
    "AI and the Future of Work",
]

INTRO_TEXT = (
    "Each week, Plumtree curates substantive research on the future of work "
    "— filtered for relevance to transformation inside life sciences and "
    "pharmaceutical R&D organizations. This brief is written for practitioners: "
    "direct, plain-language, and grounded in evidence."
)


def create_brief_document(
    brief_date: str,
    category_content: dict[str, str],
    so_what: str,
) -> bytes:
    """Create the weekly brief as a Word document and return bytes.

    Args:
        brief_date: e.g. "April 4, 2026"
        category_content: mapping of category name -> formatted text
        so_what: the synthesis paragraph
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = PLUMTREE_DARK

    # --- Title block ---
    _add_title_block(doc, brief_date)

    # --- Intro ---
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = intro.add_run(INTRO_TEXT)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = PLUMTREE_DARK

    # Separator line
    _add_separator(doc)

    # --- Category sections ---
    for i, category in enumerate(CATEGORIES, 1):
        _add_category_heading(doc, f"{i}. {category.upper()}")
        _add_separator(doc)

        content = category_content.get(category, "No substantive sources this week.")
        _add_content_paragraphs(doc, content)

        doc.add_paragraph()  # spacing

    # --- The So What This Week ---
    _add_category_heading(doc, "THE SO WHAT THIS WEEK")
    _add_separator(doc)
    so_what_para = doc.add_paragraph()
    run = so_what_para.add_run(so_what)
    run.font.size = Pt(11)
    run.font.color.rgb = PLUMTREE_DARK

    # Return as bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_title_block(doc: Document, brief_date: str) -> None:
    """Add the purple title block."""
    # "PLUMTREE SERVICES" label
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = label.add_run("PLUMTREE SERVICES")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = PLUMTREE_PURPLE

    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Future of Work Weekly Research Brief")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = PLUMTREE_PURPLE

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = date_para.add_run(f"Week Ending {brief_date}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = PLUMTREE_PURPLE


def _add_separator(doc: Document) -> None:
    """Add a purple horizontal line."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run("─" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = PLUMTREE_PURPLE


def _add_category_heading(doc: Document, text: str) -> None:
    """Add a category section heading."""
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = PLUMTREE_PURPLE


def _add_content_paragraphs(doc: Document, content: str) -> None:
    """Parse structured content text and add formatted paragraphs.

    Handles bold markers (**text**) and preserves line structure.
    """
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Handle bold markers
        parts = line.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            run = para.add_run(part)
            run.font.size = Pt(11)
            run.font.color.rgb = PLUMTREE_DARK
            if i % 2 == 1:  # odd indices are inside ** markers
                run.bold = True
